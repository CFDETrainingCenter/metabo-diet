#!/usr/bin/env python3
"""Build the Metabo-Diet Word guides, learner template bundle, and release copies.

The curriculum Markdown and assessment JSON are authoritative. This builder
applies the compact_reference_guide design tokens and the editorial_cover
opening pattern without changing the source curriculum.
"""

from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
MODULE = ROOT / "module"
SUPPORT = MODULE / "support"
QA = MODULE / "qa"

LESSONS = [
    MODULE / "content" / "lesson_01_why_harmonization_matters.md",
    MODULE / "content" / "lesson_02_comparing_study_design.md",
    MODULE / "content" / "lesson_03_harmonizing_metabolomics_metadata.md",
    MODULE / "content" / "lesson_04_guided_analysis_interpretation.md",
    MODULE / "content" / "lesson_05_access_tiers_transfer.md",
]

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
INK = "222222"
MUTED = "5A6570"
GOLD = "9A6B12"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
TABLE_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_SIDE_DXA = 120


def clean_text(value: str) -> str:
    """Normalize characters that render inconsistently in office/PDF engines."""
    return (
        value.replace("\u00a0", " ")
        .replace("\u2010", "-")
        .replace("\u2011", "-")
        .replace("\u2012", "-")
        .replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u2212", "-")
        .replace("<br>", "\n")
        .replace("<br/>", "\n")
        .replace("<br />", "\n")
    )


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (
        ("top", CELL_TOP_BOTTOM_DXA),
        ("bottom", CELL_TOP_BOTTOM_DXA),
        ("start", CELL_SIDE_DXA),
        ("end", CELL_SIDE_DXA),
    ):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def set_row_cant_split(row) -> None:
    """Keep a logical table row on one page when the row fits on a page."""
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:cantSplit"))
    if marker is None:
        marker = OxmlElement("w:cantSplit")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def set_table_geometry(table, widths: Sequence[int]) -> None:
    """Apply exact fixed-DXA geometry required by the selected preset."""
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        set_row_cant_split(row)
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field_name
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: str = INK,
    bold=None,
    italic=None,
    name="Calibri",
) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_pr.get_or_add_rFonts().set(qn("w:ascii"), name)
    r_pr.get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    """Resolve compact_reference_guide into explicit Word style tokens."""
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25
    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.0
    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(4)
    subtitle.paragraph_format.line_spacing = 1.1
    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.0
    for style_name in ("Header", "Footer"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(8.5)
        style.font.color.rgb = RGBColor.from_string(MUTED)
        style.paragraph_format.space_after = Pt(0)


def add_numbering_definitions(doc: Document) -> tuple[int, int]:
    """Add reusable bullet and decimal definitions with exact preset indents."""
    numbering = doc.part.numbering_part.element
    existing_abs = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    next_abs = max(existing_abs or [0]) + 1

    def make_abstract(abstract_id: int, kind: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        nsid = OxmlElement("w:nsid")
        nsid.set(qn("w:val"), f"{abstract_id:08X}"[-8:])
        abstract.append(nsid)
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "multilevel")
        abstract.append(multi)
        for level in range(3):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(level))
            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            lvl.append(start)
            fmt = OxmlElement("w:numFmt")
            fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
            lvl.append(fmt)
            text = OxmlElement("w:lvlText")
            text.set(qn("w:val"), "\u2022" if kind == "bullet" else f"%{level + 1}.")
            lvl.append(text)
            align = OxmlElement("w:lvlJc")
            align.set(qn("w:val"), "left")
            lvl.append(align)
            p_pr = OxmlElement("w:pPr")
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "num")
            tab.set(qn("w:pos"), str(540 + level * 360))
            tabs.append(tab)
            p_pr.append(tabs)
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(540 + level * 360))
            ind.set(qn("w:hanging"), "270")
            p_pr.append(ind)
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:after"), "80")
            spacing.set(qn("w:line"), "300")
            spacing.set(qn("w:lineRule"), "auto")
            p_pr.append(spacing)
            lvl.append(p_pr)
            if kind == "bullet":
                r_pr = OxmlElement("w:rPr")
                fonts = OxmlElement("w:rFonts")
                fonts.set(qn("w:ascii"), "Symbol")
                fonts.set(qn("w:hAnsi"), "Symbol")
                r_pr.append(fonts)
                lvl.append(r_pr)
            abstract.append(lvl)
        numbering.append(abstract)

    make_abstract(next_abs, "bullet")
    make_abstract(next_abs + 1, "decimal")
    return next_abs, next_abs + 1


def new_num_id(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    existing = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    num_id = max(existing or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    for level in range(3):
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), str(level))
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        override.append(start_override)
        num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, 2)))
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num_ref))
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_hyperlink(paragraph, text: str, url: str, *, bold=False, italic=False) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    r_pr.extend((fonts, color, underline, size))
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if italic:
        r_pr.append(OxmlElement("w:i"))
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = clean_text(text)
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(
    r"(\*\*.+?\*\*|\*[^*]+\*|\[[^\]]+\]\([^)]+\)|"
    + re.escape(chr(96))
    + r".+?"
    + re.escape(chr(96))
    + r")"
)


def add_inline(paragraph, text: str, *, base_size: float | None = None, color: str = INK) -> None:
    text = clean_text(text)
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, size=base_size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=color, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size, color=color, italic=True)
        elif token.startswith("["):
            matched = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if matched:
                add_hyperlink(paragraph, matched.group(1), matched.group(2))
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=(base_size or 11) - 0.5, color=DARK_BLUE, name="Consolas")
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "EEF1F4")
            run._element.get_or_add_rPr().append(shd)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=base_size, color=color)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    style: str | None = None,
    align=None,
    before=None,
    after=None,
) -> object:
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(before)
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
    add_inline(paragraph, text)
    return paragraph


def table_widths(rows: Sequence[Sequence[str]]) -> list[int]:
    count = len(rows[0])
    maxima = []
    for index in range(count):
        values = [clean_text(row[index] if index < len(row) else "") for row in rows]
        maxima.append(
            max(6, min(54, max((len(line) for value in values for line in value.splitlines()), default=6)))
        )
    # Reserve a readable floor before distributing the remaining width.  The
    # earlier proportional rescale could shrink four-column label cells below
    # one inch, causing words such as "Phenotype" to break mid-word.
    if count == 1:
        minima = [TABLE_DXA]
    elif count == 2:
        minima = [1800, 1800]
    elif count == 3:
        minima = [1500, 1500, 1500]
    elif count == 4:
        minima = [1300, 1400, 1400, 1400]
    elif count == 5:
        minima = [900] * count
    else:
        minima = [min(600, TABLE_DXA // count)] * count
    remaining = max(0, TABLE_DXA - sum(minima))
    total_weight = sum(maxima)
    widths = [
        minimum + int(remaining * value / total_weight)
        for minimum, value in zip(minima, maxima)
    ]
    widths[-1] += TABLE_DXA - sum(widths)
    return widths


def add_table(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    normalized = [list(row) + [""] * (column_count - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=column_count)
    table.style = "Table Grid"
    widths = table_widths(normalized)
    set_table_geometry(table, widths)
    font_size = 9.2 if column_count <= 3 else 8.3 if column_count <= 5 else 7.5
    for row_index, row in enumerate(normalized):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.12
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if len(value) <= 18 and col_index > 0
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            add_inline(paragraph, value, base_size=font_size)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
        if row_index == 0:
            set_repeat_table_header(table.rows[row_index])
    tail = doc.add_paragraph()
    tail.paragraph_format.space_before = Pt(0)
    tail.paragraph_format.space_after = Pt(2)


def parse_table_line(line: str) -> list[str]:
    return [clean_text(part.strip()) for part in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?-{3,}", line)) and "|" in line


def set_paragraph_shading(paragraph, fill: str, border_color: str = BLUE) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_color)
    borders.append(left)
    p_pr.append(borders)


def add_markdown(
    doc: Document,
    path: Path,
    *,
    skip_h1=False,
    explicit_heading_separation=False,
    page_break_before=False,
) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    compact_lesson_sources = path.name in {
        "lesson_02_comparing_study_design.md",
        "lesson_04_guided_analysis_interpretation.md",
        "lesson_05_access_tiers_transfer.md",
    }
    in_compact_source_list = False
    pending_page_break = page_break_before
    bullet_abs, decimal_abs = doc._metabo_number_abs
    current_kind = None
    current_num_id = None
    in_code = False
    in_multiple_choice_question = False
    index = 0
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if stripped.startswith(chr(96) * 3):
            in_code = not in_code
            current_kind = None
            current_num_id = None
            index += 1
            continue
        if in_code:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.right_indent = Inches(0.05)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            set_paragraph_shading(paragraph, "F3F5F7", "B8C3CC")
            run = paragraph.add_run(clean_text(raw) if raw else " ")
            set_run_font(run, size=8.1, color="24313B", name="Consolas")
            index += 1
            continue
        if not stripped or stripped == "---":
            current_kind = None
            current_num_id = None
            index += 1
            continue
        if (
            stripped.startswith("|")
            and index + 1 < len(lines)
            and is_table_separator(lines[index + 1])
        ):
            rows = [parse_table_line(raw)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(parse_table_line(lines[index]))
                index += 1
            add_table(doc, rows)
            current_kind = None
            current_num_id = None
            in_multiple_choice_question = False
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            if skip_h1 and level == 1:
                index += 1
                continue
            level = min(3, level)
            if explicit_heading_separation and index > 0:
                # A real paragraph boundary closes any preceding list for PDF
                # reading order and keeps attribution headings from being
                # concatenated with the prior list item or prose by extractors.
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(0)
                spacer.paragraph_format.space_after = Pt(1)
                spacer.paragraph_format.keep_with_next = True
            heading_text = heading.group(2)
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            if pending_page_break:
                paragraph.paragraph_format.page_break_before = True
                pending_page_break = False
            add_inline(paragraph, heading_text)
            in_compact_source_list = bool(
                compact_lesson_sources
                and heading_text.casefold() == "primary sources and first-party documentation"
            )
            current_kind = None
            current_num_id = None
            in_multiple_choice_question = False
            index += 1
            continue
        quote = re.match(r"^>\s?(.*)$", stripped)
        if quote:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.right_indent = Inches(0.08)
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(6)
            set_paragraph_shading(paragraph, LIGHT_GRAY)
            add_inline(paragraph, quote.group(1), color=DARK_BLUE)
            current_kind = None
            current_num_id = None
            index += 1
            continue
        list_match = re.match(r"^(\s*)([-+*]|\d+\.)\s+(.+)$", raw)
        if list_match:
            indent = len(list_match.group(1).replace("\t", "    "))
            level = min(2, indent // 2)
            marker = list_match.group(2)
            kind = "bullet" if marker in {"-", "+", "*"} else "decimal"
            if current_kind != kind or current_num_id is None:
                current_num_id = new_num_id(doc, bullet_abs if kind == "bullet" else decimal_abs)
                current_kind = kind
            paragraph = doc.add_paragraph()
            apply_numbering(paragraph, current_num_id, level)
            if in_compact_source_list:
                # These two six-item first-party source lists otherwise leave
                # only their final citation on a nearly blank continuation
                # page immediately before the next major lesson break.
                if path.name == "lesson_05_access_tiers_transfer.md":
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                else:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.1
            if explicit_heading_separation:
                # The attribution appendix contains several short source lists.
                # Keep their rhythm compact so a final licence line is not
                # stranded at the top of the following page.
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.1
            add_inline(paragraph, list_match.group(3))
            index += 1
            continue
        paragraph = doc.add_paragraph()
        is_question_prompt = bool(re.match(r"^\*\*KC\d+-\d+\.\*\*", stripped))
        is_question_option = bool(re.match(r"^[A-D]\.\s+", stripped))
        if is_question_prompt:
            in_multiple_choice_question = True
        if in_multiple_choice_question and (is_question_prompt or is_question_option):
            # A learner should never have to turn the page between a prompt
            # and its answer choices. The complete five-paragraph block fits
            # on one page at the guide's body size.
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.keep_with_next = not stripped.startswith("D. ")
        if path.name == "glossary.md" and re.fullmatch(r"\*\*[^*]+\*\*", stripped):
            # Keep each glossary term with at least the first paragraph of its
            # definition instead of leaving the term alone at a page foot.
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.keep_with_next = True
        add_inline(paragraph, stripped)
        if in_multiple_choice_question and stripped.startswith("D. "):
            in_multiple_choice_question = False
        current_kind = None
        current_num_id = None
        index += 1


def set_page_layout(doc: Document, running_label: str) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(running_label)
        set_run_font(run, size=8.5, color=MUTED, bold=True)
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run("Metabo-Diet  |  Page ")
        set_run_font(run, size=8.5, color=MUTED)
        add_page_field(paragraph, "PAGE")
        run = paragraph.add_run(" of ")
        set_run_font(run, size=8.5, color=MUTED)
        add_page_field(paragraph, "NUMPAGES")


def make_document(title: str, subject: str, running_label: str) -> Document:
    doc = Document()
    configure_styles(doc)
    set_page_layout(doc, running_label)
    doc._metabo_number_abs = add_numbering_definitions(doc)
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "Metabo-Diet Training Module"
    props.keywords = "metabolomics, harmonization, diet, exercise, Metabolomics Workbench, CFDE"
    props.comments = "Generated from version-controlled curriculum sources."
    return doc


def add_cover(
    doc: Document,
    *,
    packet_type: str,
    subtitle: str,
    audience: str,
) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(90)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run(packet_type.upper())
    set_run_font(run, size=10.5, color=GOLD, bold=True)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Metabo-Diet")
    subtitle_p = doc.add_paragraph(style="Subtitle")
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.add_run("Harmonizing Dietary and Exercise Phenotypes")
    subtitle_p2 = doc.add_paragraph(style="Subtitle")
    subtitle_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p2.paragraph_format.space_after = Pt(28)
    subtitle_p2.add_run("with Metabolomics Across CFDE Resources")
    descriptor = doc.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_after = Pt(72)
    run = descriptor.add_run(subtitle)
    set_run_font(run, size=10.5, color=GOLD, bold=True)
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_after = Pt(4)
    run = date.add_run("Release 1.0  |  August 2026")
    set_run_font(run, size=12, color=NAVY, bold=True)
    prepared = doc.add_paragraph()
    prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared.paragraph_format.space_after = Pt(18)
    run = prepared.add_run(audience)
    set_run_font(run, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()


def add_contents(doc: Document, entries: Sequence[str]) -> None:
    doc.add_heading("Contents", level=1)
    lead = doc.add_paragraph()
    add_inline(
        lead,
        "Use Word's Navigation pane or the descriptive headings in the PDF to move through the packet.",
    )
    for number, entry in enumerate(entries, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        run = p.add_run(f"{number:02d}  ")
        set_run_font(run, size=10.5, color=GOLD, bold=True)
        run = p.add_run(entry)
        set_run_font(run, size=11, color=NAVY, bold=True)
    doc.add_page_break()


def add_assessment(
    doc: Document,
    path: Path,
    heading: str,
    *,
    include_keys: bool = False,
    page_break_before: bool = False,
) -> None:
    payload = json.loads(path.read_text(encoding="utf-8"))
    heading_paragraph = doc.add_heading(heading, level=1)
    heading_paragraph.paragraph_format.page_break_before = page_break_before
    add_paragraph(doc, payload.get("instructions", ""), after=8)
    time_minutes = payload.get("time_minutes")
    if time_minutes:
        add_paragraph(doc, f"Estimated time: {time_minutes} minutes", after=10)
    for number, item in enumerate(payload["items"], start=1):
        p = doc.add_paragraph(style="Heading 2")
        p.paragraph_format.keep_with_next = True
        add_inline(p, f"{number}. {item['prompt']}")
        for key, choice in item["options"].items():
            option = doc.add_paragraph()
            option.paragraph_format.left_indent = Inches(0.38)
            option.paragraph_format.first_line_indent = Inches(-0.22)
            option.paragraph_format.space_after = Pt(3)
            # Keep each answer choice linked to the next choice/response so a
            # learner never encounters orphaned options on the following page.
            option.paragraph_format.keep_with_next = True
            run = option.add_run(f"[ ] {key}. ")
            set_run_font(run, bold=True, color=DARK_BLUE)
            add_inline(option, choice)
        response = doc.add_paragraph()
        response.paragraph_format.space_after = Pt(8)
        run = response.add_run("Response: ________    Confidence (1-5): ________")
        set_run_font(run, size=9.5, color=MUTED, italic=True)
        if include_keys:
            key_p = doc.add_paragraph()
            set_paragraph_shading(key_p, LIGHT_GRAY, GOLD)
            add_inline(key_p, f"Answer: {item['answer']}. {item['rationale']}")


def add_provenance_appendix(doc: Document) -> None:
    provenance = json.loads((MODULE / "data" / "provenance.json").read_text(encoding="utf-8"))
    doc.add_heading("Appendix - Data provenance", level=1)
    add_paragraph(
        doc,
        "This appendix summarizes the versioned evidence trail used by the release. "
        "The machine-readable source is module/data/provenance.json.",
    )
    repository = provenance["repository"]
    add_table(
        doc,
        [
            ["Provenance field", "Release value"],
            ["Repository", repository["name"]],
            ["Retrieval time (UTC)", repository["retrieved_at_utc"]],
            ["Access model", repository["access_model"]],
            ["Selection status", provenance["selection_status"]],
            ["Manifest schema", provenance["schema_version"]],
        ],
    )
    doc.add_heading("Locked study pair", level=2)
    study_rows = [["Role", "Accession", "Specimen", "Design summary", "License"]]
    for study in provenance["studies"]:
        study_rows.append(
            [
                study["role"].title(),
                study["study_id"],
                study["specimen"],
                study["design"],
                study["release"]["license"],
            ]
        )
    add_table(doc, study_rows)
    overlap = provenance["primary_pair"]["overlap"]
    doc.add_heading("RefMet overlap audit", level=2)
    add_table(
        doc,
        [
            ["Measure", "Value"],
            ["Diet unique nonblank RefMet names", str(overlap["diet_unique_refmet_names"])],
            ["Exercise unique nonblank RefMet names", str(overlap["exercise_unique_refmet_names"])],
            ["Raw exact overlap", str(overlap["raw_exact_overlap"])],
            [
                "Known isotope/internal-standard exclusions",
                str(len(overlap["excluded_internal_standard_mappings"])),
            ],
            [
                "Conservative biological overlap",
                str(overlap["recommended_conservative_biological_overlap"]),
            ],
        ],
    )
    add_paragraph(doc, overlap["interpretation"])
    doc.add_heading("Sanitization and claim boundaries", level=2)
    for group, rules in provenance["sanitization_rules"].items():
        if isinstance(rules, str):
            add_paragraph(doc, rules)
            continue
        p = doc.add_paragraph(style="Heading 3")
        add_inline(p, group.replace("_", " ").title())
        num_id = new_num_id(doc, doc._metabo_number_abs[0])
        for rule in rules:
            item = doc.add_paragraph()
            apply_numbering(item, num_id, 0)
            add_inline(item, rule)
    doc.add_heading("Cached-source integrity", level=2)
    rows = [["Study", "Cached file", "Bytes", "SHA-256"]]
    for study in provenance["studies"]:
        for cached in study.get("cached_files", []):
            rows.append(
                [
                    study["study_id"],
                    Path(cached["path"]).name,
                    str(cached["bytes"]),
                    cached["sha256"],
                ]
            )
    refmet = provenance.get("refmet_cache", {})
    if refmet:
        rows.append(
            [
                "RefMet",
                Path(refmet["path"]).name,
                str(refmet["bytes"]),
                refmet["sha256"],
            ]
        )
    add_table(doc, rows)


def set_inline_shape_alt_text(shape, *, title: str, description: str) -> None:
    """Set meaningful title/description text on an inline Word image."""
    doc_properties = shape._inline.docPr
    doc_properties.set("name", clean_text(title))
    doc_properties.set("title", clean_text(title))
    doc_properties.set("descr", clean_text(description))


def add_pca_examples_appendix(doc: Document) -> None:
    """Add the two release PCA figures with captions and interpretation limits."""
    doc.add_heading("Appendix - PCA worked examples", level=1)
    add_paragraph(
        doc,
        "These figures are the exact worked examples generated by notebook Lesson 4. "
        "Find the corresponding code at NB-L4-PCA-DIET and NB-L4-PCA-EXERCISE in "
        "module/notebooks/metabo_diet_harmonization.ipynb. Each model has its own "
        "coordinate system; do not compare point positions or axis values across figures.",
    )
    figures = [
        {
            "heading": "Figure A1 - Diet-study plasma PCA",
            "path": MODULE / "figures" / "ST001521_AN002534_pca.png",
            "caption": (
                "Figure A1. ST001521 / AN002534 within-study PCA of 150 biological "
                "plasma samples. The 10 pooled QPP samples are excluded."
            ),
            "alt": (
                "Scatterplot of 150 ST001521 plasma samples in a PCA fit only to analysis "
                "AN002534. Color identifies Western, Vegan, or Modulen diet labels and marker "
                "shape identifies Baseline, Day 5, Day 9, Day 12, or Day 15. PC1 explains "
                "34.4 percent and PC2 explains 11.6 percent of variance."
            ),
            "interpretation": (
                "How to read it: describe clustering or overlap as exploratory structure. "
                "The plot does not isolate a diet effect because setting, prior diet, "
                "antibiotics, PEG, time, and other design features can contribute."
            ),
        },
        {
            "heading": "Figure A2 - Exercise-study serum PCA",
            "path": MODULE / "figures" / "ST003348_AN005483_pca.png",
            "caption": (
                "Figure A2. ST003348 / AN005483 within-study PCA of 76 serum samples "
                "after excluding explicit isotope-labeled/internal-standard rows."
            ),
            "alt": (
                "Scatterplot of 76 ST003348 serum samples in a PCA fit only to analysis "
                "AN005483. Color identifies rest, immediate post-exercise, 3-hour recovery, "
                "or 22-hour recovery. PC1 explains 16.1 percent and PC2 explains 8.4 percent "
                "of variance."
            ),
            "interpretation": (
                "How to read it: the timepoint pattern is exploratory and can also reflect "
                "fasting, clock time, repeated measures, or analytical structure. PCA alone "
                "is not a repeated-measures test and does not prove causation."
            ),
        },
    ]
    for index, figure in enumerate(figures):
        if index:
            doc.add_page_break()
        heading = doc.add_heading(figure["heading"], level=2)
        heading.paragraph_format.keep_with_next = True
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.paragraph_format.space_after = Pt(4)
        image_paragraph.paragraph_format.keep_with_next = True
        shape = image_paragraph.add_run().add_picture(str(figure["path"]), width=Inches(6.25))
        set_inline_shape_alt_text(
            shape,
            title=figure["heading"],
            description=figure["alt"],
        )
        caption = doc.add_paragraph(style="Caption")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_before = Pt(0)
        caption.paragraph_format.space_after = Pt(6)
        add_inline(caption, figure["caption"], base_size=9.5, color=MUTED)
        interpretation = doc.add_paragraph()
        set_paragraph_shading(interpretation, LIGHT_GRAY, GOLD)
        add_inline(interpretation, figure["interpretation"], color=DARK_BLUE)


def build_learner_guide(output: Path) -> None:
    doc = make_document(
        "Metabo-Diet Learner Guide",
        "Five-lesson learner guide for phenotype-to-metabolome harmonization.",
        "METABO-DIET  |  LEARNER GUIDE",
    )
    add_cover(
        doc,
        packet_type="Learner Guide",
        subtitle="A guided, self-paced module",
        audience="Prepared for CFDE learners working with public metabolomics data",
    )
    add_contents(
        doc,
        [
            "How to use this guide",
            "Before you begin - first-time setup",
            "Pretest",
            "Lesson 1 - Why harmonization matters",
            "Lesson 2 - Comparing study design",
            "Lesson 3 - Harmonizing metabolomics and metadata",
            "Lesson 4 - Guided analysis and interpretation",
            "Lesson 5 - Access tiers and transfer",
            "Posttest",
            "Learner worksheets and checklists",
            "Glossary",
            "Appendix - PCA worked examples",
            "Data provenance and sources",
        ],
    )
    doc.add_heading("How to use this guide", level=1)
    add_paragraph(
        doc,
        "The module takes approximately 153 minutes after software setup: 140 minutes of lesson "
        "activity, a 5-minute pretest, and an 8-minute posttest. First-time Python or metabolomics "
        "learners should allow another 30 to 60 minutes. Extract metabo_diet_analysis_bundle.zip "
        "and metabo_diet_templates.zip, and keep this guide plus the three worksheets open. The "
        "main notebook is module/notebooks/metabo_diet_harmonization.ipynb.",
    )
    doc.add_heading("Guide-to-notebook sequence", level=2)
    sequence = [
        "Read the matching lesson in this guide.",
        "Open the notebook section with the same NB-L1 through NB-L5 key.",
        "Run each code cell after a Run now heading, from top to bottom.",
        "At a Learner edit heading, change only the named value or response, then run the cell.",
        "Compare the output with the stated result and stop if a check fails.",
        "Write the requested answer and use the Ready to move on? note before continuing.",
    ]
    sequence_num_id = new_num_id(doc, doc._metabo_number_abs[1])
    for step in sequence:
        paragraph = doc.add_paragraph()
        apply_numbering(paragraph, sequence_num_id, 0)
        add_inline(paragraph, step)
    add_table(
        doc,
        [
            ["Guide lesson", "Notebook key", "Run or complete"],
            ["Lesson 1 - Why harmonization matters", "NB-L1", "NB-SETUP; locked configuration; scientific boundary"],
            ["Lesson 2 - Comparing study design", "NB-L2", "Endpoint, sample-role, participant, and time audits"],
            ["Lesson 3 - Harmonizing metadata/metabolites", "NB-L3", "NB-L3-CROSSWALK; labeled-standard and trace exercises"],
            ["Lesson 4 - Guided analysis", "NB-L4", "NB-L4-CLASS; NB-L4-PCA-DIET; NB-L4-PCA-EXERCISE"],
            ["Lesson 5 - Access transfer", "NB-L5", "Transfer-decision learner edit; NB-REPRO; posttest"],
        ],
    )
    add_table(
        doc,
        [
            ["Release control", "Value"],
            ["Diet case study", "ST001521 - FARMM diet-anchored plasma study"],
            ["Exercise case study", "ST003348 - race-walking exercise-anchored serum study"],
            [
                "Analysis boundary",
                "Quantitative exploration stays within study; cross-study pooling is not authorized.",
            ],
            ["Offline path", "Use the versioned cache and record the fallback in the source log."],
            [
                "Visual accessibility",
                "Every proposed figure is paired with an equivalent text description in its lesson.",
            ],
        ],
    )
    doc.add_heading("Module learning objectives", level=2)
    objectives = [
        "LO1 - Explain how study design, specimen, time, and phenotype context shape interpretation.",
        "LO2 - Judge direct, partial, and non-comparability for a stated purpose.",
        "LO3 - Build a provenance-preserving RefMet crosswalk and document mapping uncertainty.",
        "LO4 - Retrieve, validate, harmonize, and explore public MW data reproducibly.",
        "LO5 - Verify access requirements and transfer the workflow safely to other resources.",
    ]
    num_id = new_num_id(doc, doc._metabo_number_abs[0])
    for text in objectives:
        p = doc.add_paragraph()
        apply_numbering(p, num_id, 0)
        add_inline(p, text)
    doc.add_page_break()
    add_markdown(doc, MODULE / "content" / "getting_started.md")
    doc.add_page_break()
    add_assessment(doc, MODULE / "assessments" / "pretest.json", "Learner pretest")
    for lesson_index, lesson in enumerate(LESSONS):
        if lesson_index == len(LESSONS) - 1:
            # Lesson 4's compact source list can end exactly at a page edge;
            # page-break-before starts Lesson 5 cleanly without generating an
            # intervening blank page.
            add_markdown(doc, lesson, page_break_before=True)
        else:
            doc.add_page_break()
            add_markdown(doc, lesson)
    # The same boundary condition occurs after Lesson 5's source list.
    add_assessment(
        doc,
        MODULE / "assessments" / "posttest.json",
        "Learner posttest",
        page_break_before=True,
    )
    for template in (
        MODULE / "templates" / "cohort_comparison_worksheet_learner.md",
        MODULE / "templates" / "metabolite_metadata_crosswalk_learner.md",
        MODULE / "templates" / "access_tier_transfer_checklist_learner.md",
    ):
        doc.add_page_break()
        add_markdown(doc, template)
    doc.add_page_break()
    add_markdown(doc, MODULE / "content" / "glossary.md")
    doc.add_page_break()
    add_pca_examples_appendix(doc)
    doc.add_page_break()
    add_provenance_appendix(doc)
    doc.add_heading("Data attribution and reuse", level=1)
    add_markdown(
        doc,
        MODULE / "ATTRIBUTION.md",
        skip_h1=True,
        explicit_heading_separation=True,
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_instructor_packet(output: Path) -> None:
    doc = make_document(
        "Metabo-Diet Instructor Packet",
        "Facilitation, assessment, guardrail, and troubleshooting packet.",
        "METABO-DIET  |  INSTRUCTOR PACKET",
    )
    add_cover(
        doc,
        packet_type="Instructor Packet",
        subtitle="Facilitation, assessment, and scientific-safety resources",
        audience="Instructor-only: keep assessment keys hidden until learner submission",
    )
    add_contents(
        doc,
        [
            "Instructor guide and session plan",
            "Objective-instruction-assessment map",
            "Scientific guardrails and troubleshooting",
            "Answer and rationale keys",
            "Instructor cohort-comparison worksheet",
            "Instructor metabolite/metadata crosswalk",
            "Instructor access-tier transfer checklist",
            "Release provenance",
        ],
    )
    add_markdown(doc, MODULE / "content" / "instructor_guide.md")
    doc.add_page_break()
    add_markdown(doc, MODULE / "assessments" / "answer_key.md")
    for template in (
        MODULE / "templates" / "cohort_comparison_worksheet_instructor.md",
        MODULE / "templates" / "metabolite_metadata_crosswalk_instructor.md",
        MODULE / "templates" / "access_tier_transfer_checklist_instructor.md",
    ):
        doc.add_page_break()
        add_markdown(doc, template)
    doc.add_page_break()
    add_provenance_appendix(doc)
    doc.add_heading("Data attribution and reuse", level=1)
    add_markdown(
        doc,
        MODULE / "ATTRIBUTION.md",
        skip_h1=True,
        explicit_heading_separation=True,
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def write_csv(path: Path, rows: Iterable[Sequence[str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerows(rows)


def build_template_bundle(output: Path) -> None:
    staging = SUPPORT / "learner_templates"
    if staging.exists():
        shutil.rmtree(staging)
    worksheet_dir = staging / "worksheets"
    worksheet_dir.mkdir(parents=True)
    for source in (
        MODULE / "templates" / "cohort_comparison_worksheet_learner.md",
        MODULE / "templates" / "metabolite_metadata_crosswalk_learner.md",
        MODULE / "templates" / "access_tier_transfer_checklist_learner.md",
    ):
        shutil.copy2(source, worksheet_dir / source.name.replace("_learner", ""))
    write_csv(
        worksheet_dir / "cohort_comparison_worksheet.csv",
        [
            [
                "section",
                "field",
                "diet_value",
                "diet_source",
                "exercise_value",
                "exercise_source",
                "comparability_D_P_N_U",
                "intended_purpose",
                "evidence_based_justification",
                "allowed_action",
                "prohibited_inference",
                "reviewer",
                "reviewed_at_utc",
            ],
            ["provenance", "study_accession", "", "", "", "", "", "", "", "", "", "", ""],
            ["population_design", "population", "", "", "", "", "", "", "", "", "", "", ""],
            ["specimen", "specific_biological_matrix", "", "", "", "", "", "", "", "", "", "", ""],
            ["time", "timepoint_anchor_offset_state", "", "", "", "", "", "", "", "", "", "", ""],
            ["phenotype", "phenotype_or_intervention", "", "", "", "", "", "", "", "", "", "", ""],
            ["assay", "assay_platform_and_mode", "", "", "", "", "", "", "", "", "", "", "", ""],
            ["assay", "units_and_scale", "", "", "", "", "", "", "", "", "", "", ""],
            ["population_design", "subject_sample_linkage", "", "", "", "", "", "", "", "", "", "", ""],
            ["provenance", "access_license_retrieval", "", "", "", "", "", "", "", "", "", "", ""],
        ],
    )
    write_csv(
        worksheet_dir / "metabolite_metadata_crosswalk.csv",
        [
            [
                "row_id",
                "config_key",
                "study_id",
                "analysis_id",
                "source_feature_id",
                "source_metabolite_name",
                "source_refmet_name",
                "source_formula",
                "source_inchikey",
                "source_external_ids",
                "source_unit",
                "source_endpoint",
                "retrieved_at_utc",
                "lookup_string",
                "lookup_transform",
                "refmet_candidate",
                "refmet_query_url_or_bulk_source_version",
                "annotation_resolution",
                "identification_evidence",
                "mapping_status",
                "source_role_or_artifact_status",
                "decision_reason",
                "reviewer_or_rule_version",
                "reviewed_at_utc",
                "eligible_exact_name_overlap",
                "eligible_class_summary",
                "eligible_within_study_quantitative",
                "eligible_cross_study_quantitative",
                "exclusion_reason",
                "decision_log_append_only",
            ],
            [
                "M001", "", "", "", "", "", "", "NR", "NR", "", "", "", "", "",
                "none", "", "", "", "", "", "", "", "", "", "review", "review",
                "review", "no", "", "",
            ],
        ],
    )
    write_csv(
        worksheet_dir / "metadata_crosswalk.csv",
        [
            [
                "metadata_row_id",
                "config_key",
                "study_accession",
                "source_block_or_endpoint",
                "source_field",
                "source_value",
                "construct_definition",
                "harmonized_field",
                "harmonized_value",
                "transformation_rule",
                "information_lost",
                "compatibility",
                "purpose",
                "source_url",
                "retrieved_at_utc",
                "reviewer_or_rule_version",
                "decision_note",
            ],
            ["MD001", "", "", "", "", "", "Specimen matrix", "", "", "", "", "", "", "", "", "", ""],
        ],
    )
    write_csv(
        worksheet_dir / "unit_compatibility_audit.csv",
        [
            [
                "config_key",
                "analysis_id",
                "source_quantity",
                "source_unit_or_scale",
                "candidate_target",
                "dimensionally_convertible",
                "quantification_compatible",
                "action",
                "formula_or_rule",
                "reason",
                "reviewer_or_rule_version",
                "reviewed_at_utc",
            ],
            ["", "", "", "", "", "", "", "", "", "", "", ""],
        ],
    )
    write_csv(
        worksheet_dir / "refmet_overlap_audit.csv",
        [
            [
                "stage",
                "diet_count",
                "exercise_count",
                "shared_count",
                "rule_or_version",
                "exclusions_at_stage",
                "source_endpoint_or_file",
                "retrieved_at_utc",
                "reviewer",
                "decision_log_append_only",
            ],
            ["Distinct submitted labels", "", "", "NA", "", "", "", "", "", ""],
            ["Nonblank source RefMet strings", "", "", "", "", "", "", "", "", ""],
            ["Reviewed accepted mappings", "", "", "", "", "", "", "", "", ""],
            ["Eligible after artifact rules", "", "", "", "", "", "", "", "", ""],
            ["Unique exact-name keys", "", "", "", "", "", "", "", "", ""],
            ["Final reported set", "", "", "", "", "", "", "", "", ""],
        ],
    )
    write_csv(
        worksheet_dir / "access_tier_transfer_checklist.csv",
        [
            [
                "section",
                "item",
                "resource",
                "exact_dataset_release",
                "release_version_or_date",
                "data_level",
                "intended_action",
                "access_pattern",
                "first_party_evidence_url",
                "first_party_document_section",
                "checked_at_utc",
                "verification_status",
                "approved_compute",
                "approved_storage",
                "permitted_outputs",
                "scientific_compatibility",
                "required_action",
                "owner",
                "mitigation",
                "stop_condition",
                "unresolved_items",
                "decision",
            ],
            ["scope", "exact dataset and intended action", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "REVISE"],
        ],
    )
    shutil.copy2(ROOT / "LICENSE", staging / "LICENSE")
    shutil.copy2(MODULE / "ATTRIBUTION.md", staging / "DATA_ATTRIBUTION.md")
    provenance_target = staging / "module" / "data" / "provenance.json"
    provenance_target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(MODULE / "data" / "provenance.json", provenance_target)
    included = sorted(
        str(path.relative_to(staging)) for path in staging.rglob("*") if path.is_file()
    )
    readme = [
        "# Metabo-Diet learner templates",
        "",
        "Editable Markdown and CSV versions of the three learner artifacts are in worksheets/.",
        "For the runnable Python notebook, cached public inputs, pipeline, figures, and R appendix,",
        "download metabo_diet_analysis_bundle.zip alongside this worksheet bundle.",
        "",
        "Scientific boundary: do not concatenate uncalibrated peak-area matrices across the two studies.",
        "Preserve the source accession, analysis ID, specimen, timepoint, and mapping evidence.",
        "Original training materials are licensed under CC BY 4.0; see LICENSE and DATA_ATTRIBUTION.md.",
        "",
        "## Included files",
        "",
    ]
    readme.extend(f"- {name}" for name in included)
    (staging / "README.md").write_text("\n".join(readme) + "\n", encoding="utf-8")
    output.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(
        output,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
    ) as archive:
        for path in sorted(staging.rglob("*")):
            if path.is_file():
                archive.write(path, path.relative_to(staging))
    shutil.rmtree(staging)


def build_analysis_bundle(output: Path) -> None:
    """Build a runnable, cache-complete analysis bundle with its module tree intact."""
    required_files = [
        ROOT / "LICENSE",
        MODULE / "README.md",
        MODULE / "ATTRIBUTION.md",
        MODULE / "data" / "README.md",
        MODULE / "data" / "provenance.json",
        MODULE / "notebooks" / "metabo_diet_harmonization.ipynb",
        MODULE / "notebooks" / "requirements.txt",
        MODULE / "notebooks" / "requirements-dev.txt",
        MODULE / "notebooks" / "install_r_packages.R",
        MODULE / "notebooks" / "metabo_diet_R_appendix.Rmd",
        MODULE / "notebooks" / "metabo_diet_R_appendix.html",
        MODULE / "scripts" / "metabo_diet_pipeline.py",
        MODULE / "scripts" / "execute_notebook.py",
        MODULE / "scripts" / "audit_live_mw.py",
        MODULE / "scripts" / "metabo_diet_R_normalization.R",
        MODULE / "scripts" / "test_R_endpoint_normalization.R",
        MODULE / "qa" / "local_pilot_protocol.md",
        MODULE / "qa" / "live_mw_audit.json",
        SUPPORT / "metabo_diet_learner_guide.pdf",
        SUPPORT / "metabo_diet_templates.zip",
    ]
    missing = [str(path.relative_to(ROOT)) for path in required_files if not path.is_file()]
    if missing:
        raise FileNotFoundError("Analysis bundle inputs missing: " + ", ".join(missing))
    tree_files = []
    for directory in (
        MODULE / "data" / "raw",
        MODULE / "data" / "derived",
        MODULE / "figures",
        MODULE / "scripts",
        MODULE / "content",
        MODULE / "templates",
        MODULE / "assessments",
    ):
        tree_files.extend(
            path
            for path in directory.rglob("*")
            if path.is_file()
            and path.suffix != ".pyc"
            and "__pycache__" not in path.parts
            and not any(part.startswith(".") for part in path.relative_to(MODULE).parts)
        )
    sources = sorted(set(required_files + tree_files))
    readme = """# Metabo-Diet learner and analysis bundle

## Start here

1. Extract this whole archive. Keep the `module/` directory intact.
2. Open `module/support/metabo_diet_learner_guide.pdf`.
3. Extract `module/support/metabo_diet_templates.zip` and keep the three worksheets nearby.
4. Follow the first-time setup in the guide or `module/content/getting_started.md`.
5. Complete the pretest, then move through guide Lessons 1 to 5 and notebook sections `NB-L1` to `NB-L5` together.

The Python notebook is the main activity. The R companion is optional. The
validated public cache is the default, so live API access is not required after
packages are installed.

## Interactive Python path

From the directory containing `module/`, run the commands for your system.

macOS or Linux:

```bash
python3.12 --version
python3.12 -m venv .venv
./.venv/bin/python -m pip install -r module/notebooks/requirements-dev.txt
./.venv/bin/python -m jupyter lab module/notebooks/metabo_diet_harmonization.ipynb
```

Windows PowerShell:

```powershell
py -3.12 --version
py -3.12 -m venv .venv
.\\.venv\\Scripts\\python.exe -m pip install -r module\\notebooks\\requirements-dev.txt
.\\.venv\\Scripts\\python.exe -m jupyter lab module\\notebooks\\metabo_diet_harmonization.ipynb
```

Python 3.11 is also supported. The setup guide explains how to run cells and
what to do when a check fails. Outputs are written to `module/data/derived/`
and `module/figures/`.

## Optional installation smoke test

This command runs every code cell without pausing for learner edits. It checks
the installation but does not complete the course and does not overwrite the
source notebook:

```bash
./.venv/bin/python module/scripts/execute_notebook.py --output metabo_diet_smoke_test.ipynb
```

The notebook imports `module/scripts/metabo_diet_pipeline.py`; do not separate
the notebook from the rest of the module tree. See `module/data/provenance.json`
and `module/ATTRIBUTION.md` for retrieval records, checksums, and licensing.

Original Metabo-Diet training materials are licensed under CC BY 4.0; see the
root `LICENSE` file. Third-party data retain the source licenses described in
`module/ATTRIBUTION.md` and `module/data/provenance.json`.

## R appendix

The source and rendered appendix are in `module/notebooks/`. The HTML is a
viewable cached-path record. Live retrieval requires the R packages named by
the R Markdown source and network access to the repository service.

Scientific boundary: the two studies are explored separately. Do not pool
uncalibrated plasma and serum peak-area matrices.
"""
    output.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(
        output,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
    ) as archive:
        archive.writestr("README.md", readme)
        for source in sources:
            archive.write(source, source.relative_to(ROOT))

    with zipfile.ZipFile(output) as archive:
        failure = archive.testzip()
        if failure:
            raise RuntimeError("Analysis bundle CRC failure: " + failure)
        names = set(archive.namelist())
        required_members = {
            "README.md",
            "LICENSE",
            "module/notebooks/metabo_diet_harmonization.ipynb",
            "module/notebooks/install_r_packages.R",
            "module/notebooks/metabo_diet_R_appendix.Rmd",
            "module/notebooks/metabo_diet_R_appendix.html",
            "module/scripts/metabo_diet_pipeline.py",
            "module/scripts/execute_notebook.py",
            "module/scripts/audit_live_mw.py",
            "module/scripts/metabo_diet_R_normalization.R",
            "module/scripts/test_R_endpoint_normalization.R",
            "module/qa/local_pilot_protocol.md",
            "module/data/provenance.json",
            "module/qa/live_mw_audit.json",
            "module/data/raw/ST001521_data.json",
            "module/data/raw/ST003348_data.json",
            "module/data/raw/refmet_classification.json",
        }
        absent = sorted(required_members - names)
        if absent:
            raise RuntimeError("Analysis bundle members missing: " + ", ".join(absent))


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def render_docx_to_pdf(docx_path: Path) -> Path:
    """Render a DOCX to its sibling PDF with an isolated LibreOffice profile."""
    soffice = shutil.which("soffice")
    if not soffice:
        raise RuntimeError("LibreOffice/soffice is required to build release PDFs")
    pdf_path = docx_path.with_suffix(".pdf")
    with tempfile.TemporaryDirectory(prefix=f"metabo_diet_{docx_path.stem}_") as temporary:
        temporary_path = Path(temporary)
        profile = temporary_path / "lo_profile"
        output_dir = temporary_path / "output"
        profile.mkdir()
        output_dir.mkdir()
        environment = os.environ.copy()
        environment["HOME"] = str(temporary_path)
        environment["TMPDIR"] = str(temporary_path)
        result = subprocess.run(
            [
                soffice,
                "--headless",
                f"-env:UserInstallation={profile.as_uri()}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(docx_path),
            ],
            env=environment,
            check=False,
            capture_output=True,
            text=True,
            timeout=300,
        )
        rendered = output_dir / pdf_path.name
        if result.returncode != 0 or not rendered.is_file() or rendered.stat().st_size == 0:
            raise RuntimeError(
                f"Failed to render {docx_path.name}: returncode={result.returncode}; "
                f"stdout={result.stdout.strip()}; stderr={result.stderr.strip()}"
            )
        shutil.copy2(rendered, pdf_path)
    return pdf_path


def main() -> None:
    SUPPORT.mkdir(parents=True, exist_ok=True)
    QA.mkdir(parents=True, exist_ok=True)
    learner_docx = SUPPORT / "metabo_diet_learner_guide.docx"
    instructor_docx = SUPPORT / "metabo_diet_instructor_packet.docx"
    templates_zip = SUPPORT / "metabo_diet_templates.zip"
    analysis_zip = SUPPORT / "metabo_diet_analysis_bundle.zip"
    build_learner_guide(learner_docx)
    build_instructor_packet(instructor_docx)
    learner_pdf = render_docx_to_pdf(learner_docx)
    instructor_pdf = render_docx_to_pdf(instructor_docx)
    build_template_bundle(templates_zip)
    build_analysis_bundle(analysis_zip)
    report = {
        "builder": str(Path(__file__).relative_to(ROOT)),
        "design_preset": "compact_reference_guide",
        "header_pattern": "editorial_cover",
        "named_overrides": {
            "cover_title": "30 pt centered navy",
            "dense_table_text": "9.2 pt for <=3 columns; 8.3 pt for 4-5; 7.5 pt for >=6",
            "code_block_text": "8.1 pt Consolas on light gray",
        },
        "outputs": [
            {
                "path": str(path.relative_to(ROOT)),
                "bytes": path.stat().st_size,
                "sha256": sha256(path),
            }
            for path in (
                learner_docx,
                learner_pdf,
                instructor_docx,
                instructor_pdf,
                templates_zip,
                analysis_zip,
            )
        ],
    }
    (QA / "packaging_build_report.json").write_text(
        json.dumps(report, indent=2) + "\n",
        encoding="utf-8",
    )
    print(json.dumps(report, indent=2))


if __name__ == "__main__":
    main()
