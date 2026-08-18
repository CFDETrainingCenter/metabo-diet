#!/usr/bin/env python3
"""Audit final Metabo-Diet release packages and emit checksummed QA evidence."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[2]
MODULE = ROOT / "module"
SUPPORT = MODULE / "support"
SCORM = MODULE / "scorm" / "metabo_diet_scorm_1_2.zip"
QA = MODULE / "qa"
TABLE_DXA = 9360


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def record(checks: list[dict], failures: list[str], label: str, passed: bool, detail="") -> None:
    checks.append({"label": label, "passed": bool(passed), "detail": str(detail)})
    if not passed:
        failures.append(label + (": " + str(detail) if detail else ""))


def document_text(doc: Document) -> str:
    chunks = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def audit_docx(path: Path, kind: str, checks: list[dict], failures: list[str]) -> dict:
    doc = Document(path)
    text = document_text(doc)
    label = kind.replace("_", " ")
    record(checks, failures, f"{label}: nonempty title metadata", bool(doc.core_properties.title))
    record(checks, failures, f"{label}: letter page size", all(
        section.page_width == Inches(8.5) and section.page_height == Inches(11)
        for section in doc.sections
    ))
    record(checks, failures, f"{label}: one-inch margins", all(
        section.top_margin == Inches(1)
        and section.bottom_margin == Inches(1)
        and section.left_margin == Inches(1)
        and section.right_margin == Inches(1)
        for section in doc.sections
    ))
    normal = doc.styles["Normal"]
    record(
        checks,
        failures,
        f"{label}: compact-reference body type",
        normal.font.name == "Calibri" and normal.font.size == Pt(11),
        f"font={normal.font.name}, size={normal.font.size.pt if normal.font.size else None}",
    )
    heading_expectations = {
        "Heading 1": (16, "2E74B5"),
        "Heading 2": (13, "2E74B5"),
        "Heading 3": (12, "1F4D78"),
    }
    for style_name, (size, color) in heading_expectations.items():
        style = doc.styles[style_name]
        actual_color = str(style.font.color.rgb) if style.font.color and style.font.color.rgb else ""
        record(
            checks,
            failures,
            f"{label}: {style_name} preset tokens",
            style.font.name == "Calibri"
            and style.font.size == Pt(size)
            and style.font.bold is True
            and actual_color == color,
            f"font={style.font.name}, size={style.font.size.pt if style.font.size else None}, color={actual_color}",
        )
    numbered = sum(
        1
        for paragraph in doc.paragraphs
        if paragraph._p.pPr is not None and paragraph._p.pPr.find(qn("w:numPr")) is not None
    )
    record(checks, failures, f"{label}: semantic Word numbering present", numbered > 0, numbered)
    footer_xml = "\n".join(section.footer._element.xml for section in doc.sections)
    record(checks, failures, f"{label}: PAGE and NUMPAGES fields", "PAGE" in footer_xml and "NUMPAGES" in footer_xml)
    header_text = " ".join(section.header.paragraphs[0].text for section in doc.sections)
    record(checks, failures, f"{label}: running header", "METABO-DIET" in header_text)

    table_findings = []
    for table_index, table in enumerate(doc.tables, start=1):
        grid_nodes = table._tbl.tblGrid.findall(qn("w:gridCol"))
        widths = [int(node.get(qn("w:w"))) for node in grid_nodes]
        tbl_pr = table._tbl.tblPr
        layout = tbl_pr.find(qn("w:tblLayout"))
        indent = tbl_pr.find(qn("w:tblInd"))
        width_node = tbl_pr.find(qn("w:tblW"))
        row_integrity = all(
            row._tr.trPr is not None
            and row._tr.trPr.find(qn("w:cantSplit")) is not None
            for row in table.rows
        )
        header_repeat = bool(
            table.rows
            and table.rows[0]._tr.trPr is not None
            and table.rows[0]._tr.trPr.find(qn("w:tblHeader")) is not None
        )
        fixed_heights = any(
            row._tr.trPr is not None and row._tr.trPr.find(qn("w:trHeight")) is not None
            for row in table.rows
        )
        passed = (
            bool(widths)
            and sum(widths) == TABLE_DXA
            and layout is not None
            and layout.get(qn("w:type")) == "fixed"
            and indent is not None
            and indent.get(qn("w:w")) == "120"
            and width_node is not None
            and width_node.get(qn("w:w")) == str(TABLE_DXA)
            and row_integrity
            and header_repeat
            and not fixed_heights
        )
        if not passed:
            table_findings.append(
                {
                    "table": table_index,
                    "grid_sum": sum(widths),
                    "row_cant_split": row_integrity,
                    "header_repeat": header_repeat,
                    "fixed_heights": fixed_heights,
                }
            )
    record(
        checks,
        failures,
        f"{label}: exact table geometry and pagination controls",
        not table_findings and bool(doc.tables),
        table_findings,
    )

    required_common = ["ST001521", "ST003348", "153", "145", "Data provenance"]
    record(
        checks,
        failures,
        f"{label}: locked studies and overlap evidence",
        all(token.lower() in text.lower() for token in required_common),
    )
    stale_phrases = [
        "A combined PCA can visualize",
        "Why there is no inferential combined PCA",
        "research count 152",
    ]
    record(
        checks,
        failures,
        f"{label}: superseded combined-PCA and count text absent",
        not any(phrase.lower() in text.lower() for phrase in stale_phrases),
    )
    if kind == "learner_guide":
        normalized = re.sub(r"[-\s]+", " ", text.casefold())
        required_groups = [
            ("learner pretest",),
            ("learner posttest",),
            ("lesson 1",),
            ("lesson 2",),
            ("lesson 3",),
            ("lesson 4",),
            ("lesson 5",),
            ("glossary",),
            ("cohort comparison worksheet",),
            ("metabolite and metadata crosswalk",),
            ("access pattern and transfer checklist", "access tier transfer checklist"),
            ("data attribution and reuse",),
        ]
        record(
            checks,
            failures,
            "learner guide: required learner sections",
            all(any(candidate in normalized for candidate in group) for group in required_groups),
        )
        key_pattern = re.compile(r"\bAnswer:\s*[A-D]\.", re.IGNORECASE)
        record(
            checks,
            failures,
            "learner guide: pre/posttests contain no answer keys",
            "answer and rationale key" not in text.lower() and not key_pattern.search(text),
        )
        notebook_tokens = [
            "metabo_diet_harmonization.ipynb",
            "NB-L1",
            "NB-L2",
            "NB-L3",
            "NB-L4",
            "NB-L5",
            "NB-L3-CROSSWALK",
            "NB-L4-PCA-DIET",
            "NB-L4-PCA-EXERCISE",
            "Appendix - PCA worked examples",
        ]
        record(
            checks,
            failures,
            "learner guide: explicit five-lesson notebook cross-references",
            all(token in text for token in notebook_tokens),
            [token for token in notebook_tokens if token not in text],
        )
        descriptions = [
            shape._inline.docPr.get("descr", "").strip()
            for shape in doc.inline_shapes
        ]
        record(
            checks,
            failures,
            "learner guide: two inline PCA figures with meaningful alt text",
            len(doc.inline_shapes) == 2 and all(len(description) >= 80 for description in descriptions),
            {"inline_shapes": len(doc.inline_shapes), "alt_lengths": [len(value) for value in descriptions]},
        )
        record(
            checks,
            failures,
            "learner guide: PCA figure captions present",
            "Figure A1." in text and "Figure A2." in text,
        )
    else:
        normalized = re.sub(r"[-\s]+", " ", text.casefold())
        required_groups = [
            ("objective instruction assessment map", "objective instruction assessment alignment"),
            ("scientific guardrails", "scientific guardrail"),
            ("troubleshooting",),
            ("answer and rationale key", "metabo diet answer key"),
            ("cohort comparison worksheet instructor version",),
            ("metabolite and metadata crosswalk instructor version",),
            ("access pattern and transfer checklist instructor version", "access tier transfer checklist instructor version"),
        ]
        record(
            checks,
            failures,
            "instructor packet: required instructor sections",
            all(any(candidate in normalized for candidate in group) for group in required_groups),
        )
        record(
            checks,
            failures,
            "instructor packet: answer/rationale keys present",
            all(token in normalized for token in ("pretest key", "posttest key", "rationale", "pre 01", "post 12")),
        )
    with zipfile.ZipFile(path) as archive:
        archive_names = set(archive.namelist())
        document_xml = archive.read("word/document.xml").decode("utf-8")
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
        review_markup = bool(
            re.search(
                r"<w:(?:ins|del|moveFrom|moveTo|commentRangeStart)(?:\s|>)",
                document_xml,
            )
        )
        comments_part = any(name.startswith("word/comments") for name in archive_names)
        track_revisions = "<w:trackRevisions" in settings_xml
    record(
        checks,
        failures,
        f"{label}: no unresolved comments or tracked revisions",
        not review_markup and not comments_part and not track_revisions,
        {
            "review_markup": review_markup,
            "comments_part": comments_part,
            "track_revisions": track_revisions,
        },
    )
    return {
        "path": str(path.relative_to(ROOT)),
        "sha256": sha256(path),
        "bytes": path.stat().st_size,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "numbered_paragraphs": numbered,
    }


def pdfinfo(path: Path) -> dict[str, str]:
    executable = shutil.which("pdfinfo")
    if not executable:
        return {}
    result = subprocess.run(
        [executable, str(path)],
        check=False,
        capture_output=True,
        text=True,
    )
    values = {}
    for line in result.stdout.splitlines():
        if ":" in line:
            key, value = line.split(":", 1)
            values[key.strip()] = value.strip()
    return values


def audit_pdf(path: Path, expected_title: str, checks: list[dict], failures: list[str]) -> dict:
    reader = PdfReader(path)
    info = pdfinfo(path)
    pages = len(reader.pages)
    metadata_title = str((reader.metadata or {}).get("/Title", ""))
    record(checks, failures, f"{path.name}: nonempty PDF", pages > 0 and path.stat().st_size > 0)
    record(checks, failures, f"{path.name}: title metadata", metadata_title == expected_title, metadata_title)
    record(checks, failures, f"{path.name}: tagged PDF", info.get("Tagged", "").lower() == "yes", info.get("Tagged", "unavailable"))
    record(checks, failures, f"{path.name}: no PDF suspects", info.get("Suspects", "").lower() == "no", info.get("Suspects", "unavailable"))
    record(checks, failures, f"{path.name}: letter pages", "612 x 792" in info.get("Page size", ""), info.get("Page size", "unavailable"))
    text_sample = "\n".join((page.extract_text() or "") for page in reader.pages[:3])
    record(checks, failures, f"{path.name}: extractable reading text", "Metabo-Diet" in text_sample)
    return {
        "path": str(path.relative_to(ROOT)),
        "sha256": sha256(path),
        "bytes": path.stat().st_size,
        "pages": pages,
        "tagged": info.get("Tagged"),
        "suspects": info.get("Suspects"),
    }


def audit_a11y_report(path: Path, checks: list[dict], failures: list[str]) -> dict:
    payload = json.loads(path.read_text(encoding="utf-8"))
    counts = payload.get("counts", {})
    passed = all(int(counts.get(level, 0)) == 0 for level in ("high", "medium", "low"))
    record(checks, failures, f"{path.name}: zero accessibility findings", passed, counts)
    return payload


def audit_templates_zip(path: Path, checks: list[dict], failures: list[str]) -> dict:
    required = {
        "README.md",
        "LICENSE",
        "DATA_ATTRIBUTION.md",
        "module/data/provenance.json",
        "worksheets/cohort_comparison_worksheet.md",
        "worksheets/metabolite_metadata_crosswalk.md",
        "worksheets/access_tier_transfer_checklist.md",
        "worksheets/cohort_comparison_worksheet.csv",
        "worksheets/metabolite_metadata_crosswalk.csv",
        "worksheets/metadata_crosswalk.csv",
        "worksheets/unit_compatibility_audit.csv",
        "worksheets/refmet_overlap_audit.csv",
        "worksheets/access_tier_transfer_checklist.csv",
    }
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        crc = archive.testzip()
        record(checks, failures, "templates ZIP: CRC", crc is None, crc or "pass")
        record(checks, failures, "templates ZIP: learner worksheet set", required <= names, sorted(required - names))
        record(checks, failures, "templates ZIP: worksheet-focused", not any(name.endswith((".ipynb", ".Rmd", ".html")) for name in names))
        rows = list(csv.reader(archive.read("worksheets/metabolite_metadata_crosswalk.csv").decode("utf-8").splitlines()))
        header = set(rows[0]) if rows else set()
        provenance_fields = {
            "source_endpoint",
            "retrieved_at_utc",
            "refmet_query_url_or_bulk_source_version",
            "reviewed_at_utc",
            "exclusion_reason",
            "decision_log_append_only",
        }
        record(checks, failures, "templates ZIP: provenance and append-only decision-log CSV fields", provenance_fields <= header, sorted(provenance_fields - header))
        readme = archive.read("README.md").decode("utf-8")
        record(checks, failures, "templates ZIP: points to runnable analysis bundle", "metabo_diet_analysis_bundle.zip" in readme)
        return {
            "path": str(path.relative_to(ROOT)),
            "sha256": sha256(path),
            "bytes": path.stat().st_size,
            "file_count": len(names),
            "files": sorted(names),
        }


def inspect_notebook_bytes(content: bytes) -> dict:
    notebook = json.loads(content.decode("utf-8"))
    cells = notebook.get("cells", [])
    code = [cell for cell in cells if cell.get("cell_type") == "code"]
    errors = [
        output
        for cell in code
        for output in cell.get("outputs", [])
        if output.get("output_type") == "error"
    ]
    return {
        "cells": len(cells),
        "code_cells": len(code),
        "execution_counts": [cell.get("execution_count") for cell in code],
        "error_outputs": len(errors),
        "cell_ids": [cell.get("id") for cell in cells],
    }


def audit_analysis_zip(
    path: Path,
    checks: list[dict],
    failures: list[str],
    execute: bool,
    python_executable: str,
) -> dict:
    required = {
        "README.md",
        "LICENSE",
        "module/README.md",
        "module/ATTRIBUTION.md",
        "module/notebooks/metabo_diet_harmonization.ipynb",
        "module/notebooks/requirements.txt",
        "module/notebooks/requirements-dev.txt",
        "module/notebooks/install_r_packages.R",
        "module/notebooks/metabo_diet_R_appendix.Rmd",
        "module/notebooks/metabo_diet_R_appendix.html",
        "module/content/getting_started.md",
        "module/support/metabo_diet_learner_guide.pdf",
        "module/support/metabo_diet_templates.zip",
        "module/scripts/metabo_diet_pipeline.py",
        "module/scripts/execute_notebook.py",
        "module/scripts/audit_live_mw.py",
        "module/scripts/metabo_diet_R_normalization.R",
        "module/scripts/test_R_endpoint_normalization.R",
        "module/research/study_selection.md",
        "module/data/README.md",
        "module/data/provenance.json",
        "module/qa/validate_module.py",
        "module/qa/live_mw_audit.json",
        "module/qa/local_pilot_protocol.md",
        "module/data/raw/ST001521_data.json",
        "module/data/raw/ST003348_data.json",
        "module/data/raw/refmet_classification.json",
        "module/figures/ST001521_AN002534_pca.png",
        "module/figures/ST003348_AN005483_pca.png",
        "module/figures/refmet_class_summary.png",
        "module/figures/refmet_overlap_summary.png",
    }
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        crc = archive.testzip()
        record(checks, failures, "analysis ZIP: CRC", crc is None, crc or "pass")
        record(checks, failures, "analysis ZIP: runnable module tree", required <= names, sorted(required - names))
        record(checks, failures, "analysis ZIP: no absolute/traversal paths", all(not name.startswith("/") and ".." not in Path(name).parts for name in names))
        readme = archive.read("README.md").decode("utf-8")
        record(checks, failures, "analysis ZIP: deterministic clean-run command", "requirements-dev.txt" in readme and "execute_notebook.py" in readme)
        record(
            checks,
            failures,
            "analysis ZIP: beginner start sequence names guide, templates, and notebook",
            all(
                token in readme
                for token in (
                    "metabo_diet_learner_guide.pdf",
                    "metabo_diet_templates.zip",
                    "metabo_diet_harmonization.ipynb",
                )
            ),
        )
        notebook = inspect_notebook_bytes(archive.read("module/notebooks/metabo_diet_harmonization.ipynb"))
        required_cell_ids = {
            "nb-setup",
            "nb-l1",
            "nb-l2",
            "nb-l3",
            "nb-l4",
            "nb-l5",
            "nb-l3-crosswalk",
            "nb-l4-pca-diet",
            "nb-l4-pca-exercise",
            "nb-repro",
        }
        record(
            checks,
            failures,
            "analysis ZIP: saved notebook has five-lesson structure and stable IDs",
            notebook["cells"] >= 45
            and notebook["code_cells"] >= 18
            and required_cell_ids <= set(notebook["cell_ids"]),
            notebook,
        )
        record(
            checks,
            failures,
            "analysis ZIP: saved notebook has sequential execution evidence",
            notebook["execution_counts"] == list(range(1, notebook["code_cells"] + 1)),
            notebook["execution_counts"],
        )
        record(checks, failures, "analysis ZIP: saved notebook has no error outputs", notebook["error_outputs"] == 0, notebook["error_outputs"])

    execution = {"requested": execute, "passed": False}
    if execute:
        started = time.monotonic()
        with tempfile.TemporaryDirectory(prefix="metabo_diet_bundle_audit_") as temporary:
            destination = Path(temporary)
            with zipfile.ZipFile(path) as archive:
                archive.extractall(destination)
            environment = os.environ.copy()
            environment["METABO_DIET_LIVE"] = "0"
            environment["MPLCONFIGDIR"] = str(destination / "mplconfig")
            executed_notebook = destination / "metabo_diet_bundle_smoke_test.ipynb"
            command = [
                python_executable,
                str(destination / "module" / "scripts" / "execute_notebook.py"),
                "--output",
                str(executed_notebook),
            ]
            result = subprocess.run(
                command,
                cwd=destination,
                env=environment,
                check=False,
                capture_output=True,
                text=True,
                timeout=1200,
            )
            validator_result = subprocess.run(
                [
                    python_executable,
                    str(destination / "module" / "qa" / "validate_module.py"),
                ],
                cwd=destination,
                env=environment,
                check=False,
                capture_output=True,
                text=True,
                timeout=300,
            )
            output_created = executed_notebook.is_file()
            rerun = (
                inspect_notebook_bytes(executed_notebook.read_bytes())
                if output_created
                else {
                    "cells": 0,
                    "code_cells": 0,
                    "execution_counts": [],
                    "error_outputs": 0,
                    "cell_ids": [],
                }
            )
            normalized_stdout = result.stdout
            normalized_stderr = result.stderr
            normalized_validator_stdout = validator_result.stdout
            normalized_validator_stderr = validator_result.stderr
            # macOS may expose the same temporary directory through both
            # /var/... and /private/var/...; remove either spelling from QA evidence.
            temporary_paths = sorted(
                {str(destination), str(destination.resolve())}, key=len, reverse=True
            )
            for temporary_path in temporary_paths:
                normalized_stdout = normalized_stdout.replace(
                    temporary_path, "<temporary-directory>"
                )
                normalized_stderr = normalized_stderr.replace(
                    temporary_path, "<temporary-directory>"
                )
                normalized_validator_stdout = normalized_validator_stdout.replace(
                    temporary_path, "<temporary-directory>"
                )
                normalized_validator_stderr = normalized_validator_stderr.replace(
                    temporary_path, "<temporary-directory>"
                )
            normalized_stdout = normalized_stdout.strip()
            normalized_stderr = normalized_stderr.strip()
            normalized_validator_stdout = normalized_validator_stdout.strip()
            normalized_validator_stderr = normalized_validator_stderr.strip()
            execution = {
                "requested": True,
                "passed": (
                    result.returncode == 0
                    and output_created
                    and rerun["error_outputs"] == 0
                    and validator_result.returncode == 0
                ),
                "python_executable": Path(python_executable).name,
                "returncode": result.returncode,
                "elapsed_seconds": round(time.monotonic() - started, 3),
                "stdout": normalized_stdout,
                "stderr_tail": normalized_stderr[-2000:],
                "output_created": output_created,
                "notebook": rerun,
                "bundled_validator": {
                    "returncode": validator_result.returncode,
                    "stdout_tail": normalized_validator_stdout[-2000:],
                    "stderr_tail": normalized_validator_stderr[-2000:],
                },
            }
        record(
            checks,
            failures,
            "analysis ZIP: clean extraction cached Python execution and bundled validation",
            execution["passed"],
            execution,
        )
    else:
        record(
            checks,
            failures,
            "analysis ZIP: clean extraction cached Python execution and bundled validation",
            False,
            "rerun with --execute-analysis-bundle",
        )
    return {
        "path": str(path.relative_to(ROOT)),
        "sha256": sha256(path),
        "bytes": path.stat().st_size,
        "file_count": len(names),
        "notebook": notebook,
        "clean_execution": execution,
    }


def audit_scorm(path: Path, checks: list[dict], failures: list[str]) -> dict:
    validation_path = QA / "scorm_validation.json"
    validation = json.loads(validation_path.read_text(encoding="utf-8"))
    record(checks, failures, "SCORM: builder validation passed", validation.get("passed") is True, validation.get("failures"))
    record(checks, failures, "SCORM: validation hash matches archive", validation.get("sha256") == sha256(path))
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        record(checks, failures, "SCORM: root imsmanifest.xml", "imsmanifest.xml" in names)
        record(checks, failures, "SCORM: archive CRC", archive.testzip() is None)
        required_support_files = {
            "LICENSE",
            "ATTRIBUTION.md",
            "module/data/provenance.json",
        }
        record(
            checks,
            failures,
            "SCORM: license, attribution, and provenance are included",
            required_support_files <= names,
            sorted(required_support_files - names),
        )
        root = ET.fromstring(archive.read("imsmanifest.xml"))
        namespace = {"imscp": "http://www.imsproject.org/xsd/imscp_rootv1p1p2"}
        hrefs = {node.attrib.get("href", "") for node in root.findall(".//imscp:file", namespace)}
        record(checks, failures, "SCORM: manifest enumerates every payload", hrefs == names - {"imsmanifest.xml"}, sorted((names - {"imsmanifest.xml"}) ^ hrefs))
        course_js = archive.read("assets/course.js").decode("utf-8")
        lms_tokens = ["LMSInitialize", "LMSSetValue", "LMSCommit", "LMSFinish", "cmi.core.score.raw", "cmi.core.lesson_status"]
        record(checks, failures, "SCORM: LMS completion and score calls", all(token in course_js for token in lms_tokens))
        record(checks, failures, "SCORM: local progress fallback", "localStorage" in course_js and "Local progress fallback active" in course_js)
        record(checks, failures, "SCORM: bounded opener API traversal", "searchApiChain(window.opener)" in course_js and "return findApi(window.opener)" not in course_js)
        synchronized = {}
        for name in (
            "metabo_diet_learner_guide.pdf",
            "metabo_diet_templates.zip",
            "metabo_diet_analysis_bundle.zip",
        ):
            member = "downloads/" + name
            source = SUPPORT / name
            same = member in names and archive.read(member) == source.read_bytes()
            synchronized[name] = same
        record(checks, failures, "SCORM: embedded learner downloads are byte-synchronized", all(synchronized.values()), synchronized)
    return {
        "path": str(path.relative_to(ROOT)),
        "sha256": sha256(path),
        "bytes": path.stat().st_size,
        "validation": validation,
    }


def write_manifest(paths: list[Path], render_counts: dict) -> dict:
    entries = [
        {
            "path": str(path.relative_to(ROOT)),
            "bytes": path.stat().st_size,
            "sha256": sha256(path),
        }
        for path in sorted(set(paths))
    ]
    payload = {
        "release": "Metabo-Diet 1.0",
        "generated_at_utc": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
        "render_page_counts": render_counts,
        "files": entries,
    }
    (QA / "release_manifest.json").write_text(json.dumps(payload, indent=2) + "\n", encoding="utf-8")
    checksum_lines = [f"{entry['sha256']}  {entry['path']}" for entry in entries]
    (QA / "MANIFEST.sha256").write_text("\n".join(checksum_lines) + "\n", encoding="utf-8")
    return payload


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--visual-inspected", action="store_true", help="Confirm every final rendered page was manually inspected.")
    parser.add_argument("--execute-analysis-bundle", action="store_true", help="Extract and execute the cached Python notebook in a temporary directory.")
    parser.add_argument(
        "--analysis-python",
        default=sys.executable,
        help="Python interpreter containing the pinned notebook execution requirements.",
    )
    args = parser.parse_args()
    QA.mkdir(parents=True, exist_ok=True)
    checks: list[dict] = []
    failures: list[str] = []

    release_support = [
        SUPPORT / "metabo_diet_learner_guide.docx",
        SUPPORT / "metabo_diet_learner_guide.pdf",
        SUPPORT / "metabo_diet_instructor_packet.docx",
        SUPPORT / "metabo_diet_instructor_packet.pdf",
        SUPPORT / "metabo_diet_templates.zip",
        SUPPORT / "metabo_diet_analysis_bundle.zip",
    ]
    expected = sorted(set([*release_support, SCORM]))
    for path in expected:
        record(checks, failures, f"release file exists: {path.relative_to(ROOT)}", path.is_file() and path.stat().st_size > 0)
    if failures:
        report = {"passed": False, "failures": failures, "checks": checks}
        (QA / "release_packaging_audit.json").write_text(json.dumps(report, indent=2) + "\n", encoding="utf-8")
        return 1

    docx = [
        audit_docx(SUPPORT / "metabo_diet_learner_guide.docx", "learner_guide", checks, failures),
        audit_docx(SUPPORT / "metabo_diet_instructor_packet.docx", "instructor_packet", checks, failures),
    ]
    pdfs = [
        audit_pdf(SUPPORT / "metabo_diet_learner_guide.pdf", "Metabo-Diet Learner Guide", checks, failures),
        audit_pdf(SUPPORT / "metabo_diet_instructor_packet.pdf", "Metabo-Diet Instructor Packet", checks, failures),
    ]
    a11y = [
        audit_a11y_report(QA / "learner_guide_a11y.json", checks, failures),
        audit_a11y_report(QA / "instructor_packet_a11y.json", checks, failures),
    ]
    templates = audit_templates_zip(SUPPORT / "metabo_diet_templates.zip", checks, failures)
    analysis = audit_analysis_zip(
        SUPPORT / "metabo_diet_analysis_bundle.zip",
        checks,
        failures,
        args.execute_analysis_bundle,
        args.analysis_python,
    )
    scorm = audit_scorm(SCORM, checks, failures)

    byte_sync = {}
    with zipfile.ZipFile(SCORM) as archive:
        for support_path in (
            SUPPORT / "metabo_diet_learner_guide.pdf",
            SUPPORT / "metabo_diet_templates.zip",
            SUPPORT / "metabo_diet_analysis_bundle.zip",
        ):
            member = f"downloads/{support_path.name}"
            same = member in archive.namelist() and archive.read(member) == support_path.read_bytes()
            byte_sync[support_path.name] = same
    record(checks, failures, "support files and SCORM downloads are byte-synchronized", all(byte_sync.values()), byte_sync)

    render_counts = {}
    for key, pdf in zip(("learner_guide", "instructor_packet"), pdfs):
        render_dir = QA / "renders" / key
        page_pngs = sorted(render_dir.glob("page-*.png"))
        render_counts[key] = {"pdf_pages": pdf["pages"], "png_pages": len(page_pngs)}
        record(checks, failures, f"{key}: one final PNG per PDF page", len(page_pngs) == pdf["pages"], render_counts[key])
    record(
        checks,
        failures,
        "all final rendered pages manually inspected",
        args.visual_inspected,
        {
            "method": "manual full-page visual inspection of final PNG renders",
            "pages": sum(item["png_pages"] for item in render_counts.values()),
            "focused_rechecks": [
                "four-column label widths",
                "cached-source checksum row pagination",
                "attribution heading/list separation",
            ],
        },
    )

    manifest_paths = [
        ROOT / "LICENSE",
        *release_support,
        SCORM,
        QA / "packaging_build_report.json",
        QA / "scorm_validation.json",
        QA / "learner_guide_a11y.json",
        QA / "instructor_packet_a11y.json",
    ]
    manifest = write_manifest(manifest_paths, render_counts)
    report = {
        "passed": not failures,
        "failures": failures,
        "checks": checks,
        "documents": docx,
        "pdfs": pdfs,
        "accessibility": a11y,
        "templates_zip": templates,
        "analysis_bundle": analysis,
        "scorm": scorm,
        "byte_sync": byte_sync,
        "visual_inspection": {
            "complete": args.visual_inspected,
            "render_counts": render_counts,
        },
        "manifest": {
            "path": "module/qa/release_manifest.json",
            "sha256_path": "module/qa/MANIFEST.sha256",
            "file_count": len(manifest["files"]),
        },
    }
    (QA / "release_packaging_audit.json").write_text(json.dumps(report, indent=2) + "\n", encoding="utf-8")
    summary = [
        "# Metabo-Diet release packaging audit",
        "",
        f"**Result:** {'PASS' if report['passed'] else 'FAIL'}",
        f"**Learner render:** {render_counts['learner_guide']['pdf_pages']} pages",
        f"**Instructor render:** {render_counts['instructor_packet']['pdf_pages']} pages",
        f"**Visual inspection:** {'complete' if args.visual_inspected else 'not confirmed'}",
        f"**Clean extracted Python execution:** {'pass' if analysis['clean_execution']['passed'] else 'fail/not run'}",
        f"**SCORM validation:** {'pass' if scorm['validation'].get('passed') else 'fail'}",
        f"**Accessibility findings:** learner={a11y[0]['counts']}; instructor={a11y[1]['counts']}",
        "",
        "## Failures",
        "",
    ]
    summary.extend(["- None"] if not failures else [f"- {failure}" for failure in failures])
    (QA / "release_packaging_audit.md").write_text("\n".join(summary) + "\n", encoding="utf-8")
    print(json.dumps({"passed": report["passed"], "failures": failures, "render_counts": render_counts}, indent=2))
    return 0 if report["passed"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
